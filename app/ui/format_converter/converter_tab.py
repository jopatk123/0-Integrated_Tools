# -*- coding: utf-8 -*-
"""格式转换功能模块"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
from pathlib import Path
import threading
import re

try:
    import docx
    from docx import Document
    from docx.shared import Inches
except ImportError:
    docx = None

try:
    import markdown
    from markdown.extensions import tables, codehilite, toc
except ImportError:
    markdown = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

class ConverterTab:
    """格式转换功能选项卡"""
    
    def __init__(self, parent_frame, theme):
        self.parent_frame = parent_frame
        self.theme = theme
        self.update_status = None  # 状态更新回调函数
        
        # 检查依赖库
        self.check_dependencies()
        
        # 创建界面组件
        self.create_widgets()
        
    def check_dependencies(self):
        """检查必要的依赖库是否已安装"""
        missing_deps = []
        
        if docx is None:
            missing_deps.append("python-docx")
        if markdown is None:
            missing_deps.append("markdown")
        if BeautifulSoup is None:
            missing_deps.append("beautifulsoup4")
            
        if missing_deps:
            self.show_dependency_warning(missing_deps)
            
    def show_dependency_warning(self, missing_deps):
        """显示缺少依赖库的警告"""
        deps_str = ", ".join(missing_deps)
        warning_msg = f"缺少以下依赖库：{deps_str}\n\n请运行以下命令安装：\npip install {' '.join(missing_deps)}"
        
        warning_frame = tk.Frame(self.parent_frame, bg="#ffeeee", relief=tk.RAISED, bd=2)
        warning_frame.pack(fill=tk.X, padx=10, pady=5)
        
        warning_label = tk.Label(warning_frame, text=warning_msg, 
                               bg="#ffeeee", fg="#cc0000", 
                               justify=tk.LEFT, font=("微软雅黑", 10))
        warning_label.pack(padx=10, pady=10)
        
    def create_widgets(self):
        """创建界面组件"""
        # 主要操作框架
        main_frame = tk.Frame(self.parent_frame, bg=self.theme.bg_color)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 源文件选择区域
        source_frame = tk.LabelFrame(main_frame, text="源文件选择", bg=self.theme.bg_color,
                                   font=("微软雅黑", 10, "bold"))
        source_frame.pack(fill=tk.X, pady=5)
        
        # 源文件路径
        self.source_path_var = tk.StringVar()
        source_path_frame = tk.Frame(source_frame, bg=self.theme.bg_color)
        source_path_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(source_path_frame, text="源文件：", bg=self.theme.bg_color,
                font=("微软雅黑", 10)).pack(side=tk.LEFT)
        source_entry = tk.Entry(source_path_frame, textvariable=self.source_path_var, width=60,
                               font=("微软雅黑", 10))
        source_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        button_style = self.theme.get_button_style()
        browse_btn = tk.Button(source_path_frame, text="浏览", command=self.browse_source_file,
                             bg=button_style["bg"], fg=button_style["fg"],
                             font=("微软雅黑", 10))
        browse_btn.pack(side=tk.RIGHT, padx=5)
        
        # 转换选项区域
        option_frame = tk.LabelFrame(main_frame, text="转换选项", bg=self.theme.bg_color,
                                   font=("微软雅黑", 10, "bold"))
        option_frame.pack(fill=tk.X, pady=5)
        
        # 转换方向选择
        direction_frame = tk.Frame(option_frame, bg=self.theme.bg_color)
        direction_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(direction_frame, text="转换方向：", bg=self.theme.bg_color,
                font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        self.conversion_direction = tk.StringVar(value="md_to_docx")
        md_to_docx_radio = tk.Radiobutton(direction_frame, text="Markdown → Word", 
                                        variable=self.conversion_direction, value="md_to_docx",
                                        bg=self.theme.bg_color, command=self.update_file_filter,
                                        font=("微软雅黑", 10))
        md_to_docx_radio.pack(side=tk.LEFT, padx=10)
        
        docx_to_md_radio = tk.Radiobutton(direction_frame, text="Word → Markdown", 
                                        variable=self.conversion_direction, value="docx_to_md",
                                        bg=self.theme.bg_color, command=self.update_file_filter,
                                        font=("微软雅黑", 10))
        docx_to_md_radio.pack(side=tk.LEFT, padx=10)
        
        # 输出目录选择
        output_frame = tk.LabelFrame(main_frame, text="输出设置", bg=self.theme.bg_color,
                                   font=("微软雅黑", 10, "bold"))
        output_frame.pack(fill=tk.X, pady=5)
        
        self.output_dir_var = tk.StringVar()
        output_dir_frame = tk.Frame(output_frame, bg=self.theme.bg_color)
        output_dir_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(output_dir_frame, text="输出目录：", bg=self.theme.bg_color,
                font=("微软雅黑", 10)).pack(side=tk.LEFT)
        output_entry = tk.Entry(output_dir_frame, textvariable=self.output_dir_var, width=60,
                              font=("微软雅黑", 10))
        output_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        browse_output_btn = tk.Button(output_dir_frame, text="浏览", command=self.browse_output_dir,
                                    bg=button_style["bg"], fg=button_style["fg"],
                                    font=("微软雅黑", 10))
        browse_output_btn.pack(side=tk.RIGHT, padx=5)
        
        # 转换按钮
        convert_frame = tk.Frame(main_frame, bg=self.theme.bg_color)
        convert_frame.pack(fill=tk.X, pady=10)
        
        success_style = self.theme.get_button_style("success")
        self.convert_btn = tk.Button(convert_frame, text="开始转换", command=self.start_conversion,
                                   bg=success_style["bg"], fg=success_style["fg"], 
                                   font=("微软雅黑", 12, "bold"))
        self.convert_btn.pack(pady=10)
        
        # 预览按钮
        preview_btn = tk.Button(convert_frame, text="预览转换", command=self.preview_conversion,
                              bg=button_style["bg"], fg=button_style["fg"],
                              font=("微软雅黑", 10))
        preview_btn.pack(pady=5)
        
    def update_file_filter(self):
        """根据转换方向更新文件过滤器"""
        # 清空当前选择的文件
        self.source_path_var.set("")
        
    def browse_source_file(self):
        """浏览选择源文件"""
        direction = self.conversion_direction.get()
        
        if direction == "md_to_docx":
            filetypes = [("Markdown文件", "*.md *.markdown"), ("所有文件", "*.*")]
        else:
            filetypes = [("Word文档", "*.docx *.doc"), ("所有文件", "*.*")]
            
        filename = filedialog.askopenfilename(
            title="选择源文件",
            filetypes=filetypes
        )
        
        if filename:
            self.source_path_var.set(filename)
            # 自动设置输出目录为源文件所在目录
            if not self.output_dir_var.get():
                self.output_dir_var.set(os.path.dirname(filename))
                
    def browse_output_dir(self):
        """浏览选择输出目录"""
        directory = filedialog.askdirectory(title="选择输出目录")
        if directory:
            self.output_dir_var.set(directory)
            
    def start_conversion(self):
        """开始转换过程"""
        # 验证输入
        if not self.validate_inputs():
            return
            
        # 禁用转换按钮
        self.convert_btn.config(state="disabled")
        if self.update_status:
            self.update_status("转换中...")
        
        # 在新线程中执行转换
        thread = threading.Thread(target=self.perform_conversion)
        thread.daemon = True
        thread.start()
        
    def read_file_with_encoding(self, file_path):
        """尝试使用多种编码读取文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                if self.update_status:
                    self.update_status(f"使用 {encoding} 编码成功读取文件")
                return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if self.update_status:
                    self.update_status(f"读取文件时出错: {str(e)}")
                continue
        
        # 如果所有编码都失败，抛出异常
        raise UnicodeDecodeError("无法使用任何支持的编码读取文件")
    
    def validate_inputs(self):
        """验证输入参数"""
        source_file = self.source_path_var.get().strip()
        output_dir = self.output_dir_var.get().strip()
        
        if not source_file:
            messagebox.showerror("错误", "请选择源文件")
            return False
            
        if not os.path.exists(source_file):
            messagebox.showerror("错误", "源文件不存在")
            return False
            
        if not output_dir:
            messagebox.showerror("错误", "请选择输出目录")
            return False
            
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                messagebox.showerror("错误", f"无法创建输出目录：{str(e)}")
                return False
                
        # 检查依赖库
        direction = self.conversion_direction.get()
        if direction == "md_to_docx" and docx is None:
            messagebox.showerror("错误", "缺少python-docx库，请先安装")
            return False
            
        if direction == "docx_to_md" and (docx is None or BeautifulSoup is None):
            messagebox.showerror("错误", "缺少python-docx或beautifulsoup4库，请先安装")
            return False
            
        return True
        
    def perform_conversion(self):
        """执行转换操作"""
        try:
            source_file = self.source_path_var.get().strip()
            output_dir = self.output_dir_var.get().strip()
            direction = self.conversion_direction.get()
            
            if direction == "md_to_docx":
                self.convert_md_to_docx(source_file, output_dir)
            else:
                self.convert_docx_to_md(source_file, output_dir)
                
            # 转换完成
            if self.update_status:
                self.update_status("转换完成！")
            messagebox.showinfo("成功", "文件转换完成！")
            
        except Exception as e:
            if self.update_status:
                self.update_status("转换失败")
            messagebox.showerror("错误", f"转换失败：{str(e)}")
            
        finally:
            # 重新启用转换按钮
            self.convert_btn.config(state="normal")
            
    def convert_md_to_docx(self, md_file, output_dir):
        """将Markdown文件转换为Word文档"""
        # 读取Markdown文件，尝试多种编码
        md_content = self.read_file_with_encoding(md_file)
            
        # 转换为HTML
        if markdown:
            html = markdown.markdown(md_content, extensions=['tables', 'codehilite', 'toc'])
        else:
            # 简单的Markdown到HTML转换
            html = self.simple_md_to_html(md_content)
            
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        self.setup_document_styles(doc)
        
        # 解析HTML并添加到Word文档
        if BeautifulSoup:
            soup = BeautifulSoup(html, 'html.parser')
            self.html_to_docx(soup, doc)
        else:
            # 简单的文本处理
            self.simple_text_to_docx(md_content, doc)
            
        # 保存文档
        base_name = Path(md_file).stem
        output_file = os.path.join(output_dir, f"{base_name}.docx")
        try:
            doc.save(output_file)
            if self.update_status:
                self.update_status(f"文件已保存到: {output_file}")
        except Exception as e:
            raise Exception(f"保存Word文档失败: {str(e)}")
        
    def convert_docx_to_md(self, docx_file, output_dir):
        """将Word文档转换为Markdown文件"""
        # 读取Word文档
        doc = Document(docx_file)
        
        # 转换为Markdown
        md_content = self.docx_to_markdown(doc)
        
        # 保存Markdown文件
        base_name = Path(docx_file).stem
        output_file = os.path.join(output_dir, f"{base_name}.md")
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(md_content)
            if self.update_status:
                self.update_status(f"文件已保存到: {output_file}")
        except Exception as e:
            raise Exception(f"保存文件失败: {str(e)}")
            
    def simple_md_to_html(self, md_content):
        """简单的Markdown到HTML转换"""
        lines = md_content.split('\n')
        html_lines = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                html_lines.append(f'<li>{line[2:]}</li>')
            elif line:
                html_lines.append(f'<p>{line}</p>')
            else:
                html_lines.append('<br>')
                
        return '\n'.join(html_lines)
    
    def process_inline_formatting(self, text):
        """处理行内格式化"""
        if not re:
            return text
            
        # 处理粗体
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        # 处理斜体
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        # 处理行内代码
        text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
        # 处理链接
        text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', text)
        
        return text
        
    def html_to_docx(self, soup, doc):
        """将HTML内容转换为Word文档"""
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
            text = element.get_text().strip()
            if not text:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # 标题
                level = int(element.name[1])
                heading = doc.add_heading(text, level=level)
            elif element.name == 'p':
                # 段落
                doc.add_paragraph(text)
            elif element.name == 'li':
                # 列表项
                doc.add_paragraph(text, style='List Bullet')
                
    def simple_text_to_docx(self, text, doc):
        """简单的文本到Word文档转换"""
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(line)
                
    def docx_to_markdown(self, doc):
        """将Word文档转换为Markdown"""
        md_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                md_lines.append('')
                continue
                
            # 检查段落样式
            style_name = paragraph.style.name if paragraph.style else ''
            
            if 'Heading 1' in style_name:
                md_lines.append(f'# {text}')
            elif 'Heading 2' in style_name:
                md_lines.append(f'## {text}')
            elif 'Heading 3' in style_name:
                md_lines.append(f'### {text}')
            elif 'Heading 4' in style_name:
                md_lines.append(f'#### {text}')
            elif 'Heading 5' in style_name:
                md_lines.append(f'##### {text}')
            elif 'Heading 6' in style_name:
                md_lines.append(f'###### {text}')
            elif 'List' in style_name:
                md_lines.append(f'- {text}')
            else:
                md_lines.append(text)
                
        return '\n\n'.join(md_lines)
    
    def setup_document_styles(self, doc):
        """设置Word文档样式"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '微软雅黑'
            font.size = Pt(11)
            
            # 设置标题样式
            for i in range(1, 7):
                try:
                    heading_style = doc.styles[f'Heading {i}']
                    heading_font = heading_style.font
                    heading_font.name = '微软雅黑'
                    heading_font.size = Pt(16 - i)
                    heading_font.bold = True
                except KeyError:
                    pass
            
        except ImportError:
            pass  # 如果没有相关模块，跳过样式设置
    
    def preview_conversion(self):
        """预览转换结果"""
        if not self.validate_inputs():
            return
            
        try:
            source_file = self.source_path_var.get().strip()
            direction = self.conversion_direction.get()
            
            # 创建预览窗口
            preview_window = tk.Toplevel(self.parent_frame)
            preview_window.title("转换预览")
            preview_window.geometry("800x600")
            
            # 创建文本框显示预览内容
            text_frame = tk.Frame(preview_window)
            text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("微软雅黑", 10))
            scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
            text_widget.configure(yscrollcommand=scrollbar.set)
            
            text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            # 生成预览内容
            if direction == "md_to_docx":
                with open(source_file, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                
                preview_content = "Markdown 内容预览:\n\n" + md_content[:2000]
                if len(md_content) > 2000:
                    preview_content += "\n\n... (内容过长，仅显示前2000字符)"
                    
            else:  # docx_to_md
                doc = Document(source_file)
                preview_content = "Word 文档内容预览:\n\n"
                
                char_count = 0
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        preview_content += text + "\n\n"
                        char_count += len(text)
                        if char_count > 2000:
                            preview_content += "... (内容过长，仅显示部分内容)"
                            break
            
            text_widget.insert(tk.END, preview_content)
            text_widget.config(state=tk.DISABLED)
            
            # 添加关闭按钮
            close_btn = tk.Button(preview_window, text="关闭", 
                                command=preview_window.destroy,
                                font=("微软雅黑", 10))
            close_btn.pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("预览错误", f"无法生成预览：{str(e)}")